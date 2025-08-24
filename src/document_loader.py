"""
Document loading and text extraction module
Supports PDF, DOCX, and TXT file formats
"""
import os
import logging
from typing import Optional, Dict, Any
from io import BytesIO

try:
    import PyPDF2
except ImportError:
    import pypdf2 as PyPDF2
from docx import Document
from src.config import Config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentLoader:
    """Handles loading and text extraction from various document formats"""
    
    def __init__(self):
        self.config = Config()
    
    def load_document(self, file_data: BytesIO, filename: str) -> Optional[Dict[str, Any]]:
        """
        Load and extract text from uploaded document
        
        Args:
            file_data: BytesIO object containing the file data
            filename: Original filename with extension
            
        Returns:
            Dictionary containing extracted text and metadata or None if failed
        """
        try:
            # Get file extension
            file_extension = os.path.splitext(filename.lower())[1]
            
            if file_extension not in self.config.SUPPORTED_FORMATS:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Check file size
            file_data.seek(0, 2)  # Seek to end
            file_size = file_data.tell()
            file_data.seek(0)  # Reset to beginning
            
            if file_size > self.config.MAX_FILE_SIZE_MB * 1024 * 1024:
                raise ValueError(f"File size exceeds {self.config.MAX_FILE_SIZE_MB}MB limit")
            
            # Extract text based on file type
            if file_extension == '.pdf':
                text_content = self._extract_pdf_text(file_data)
            elif file_extension == '.docx':
                text_content = self._extract_docx_text(file_data)
            elif file_extension == '.txt':
                text_content = self._extract_txt_text(file_data)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            if not text_content or not text_content.strip():
                raise ValueError("No text content found in document")
            
            return {
                'text': text_content,
                'filename': filename,
                'file_type': file_extension,
                'file_size': file_size,
                'char_count': len(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error loading document {filename}: {str(e)}")
            return None
    
    def _extract_pdf_text(self, file_data: BytesIO) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(file_data)
            text_content = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    continue
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_data: BytesIO) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_data)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def _extract_txt_text(self, file_data: BytesIO) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file_data.seek(0)
                    text_content = file_data.read().decode(encoding)
                    return text_content.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            file_data.seek(0)
            text_content = file_data.read().decode('utf-8', errors='replace')
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {str(e)}")
            raise
